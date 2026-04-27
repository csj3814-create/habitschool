from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUTPUT = DOCS / "giftishow_submission_commercial_key_package_ko.docx"
SOURCE_NOTE = DOCS / "giftishow_submission_commercial_key_package_ko.md"

SCREENS = [
    {
        "heading": "1. 사용자 자산 탭 - 해빛 마켓",
        "path": DOCS / "giftishow_submission_reward_market_asset_screen.png",
        "caption": (
            "사용자가 해빛스쿨 앱의 자산 탭에서 현재 보유 포인트를 확인하고, "
            "해빛 마켓 상품을 선택해 교환을 진행하는 화면입니다. "
            "현재 운영 정책은 포인트 기반 보상 교환이며, HBT는 별도 자산 시스템으로 유지됩니다."
        ),
    },
    {
        "heading": "2. 사용자 쿠폰 보관함 - 앱 내 직접 수령",
        "path": DOCS / "giftishow_submission_reward_market_coupon_screen.png",
        "caption": (
            "발급된 쿠폰은 문자 기본 발송 대신 앱 내 쿠폰 보관함에서 확인하도록 설계했습니다. "
            "couponImgUrl, PIN 번호, 유효기간, 발급 상태를 함께 노출해 사용자가 앱에서 바로 쿠폰을 사용할 수 있습니다."
        ),
    },
    {
        "heading": "3. 관리자 관제탑 - 보상 마켓 운영",
        "path": DOCS / "giftishow_submission_reward_market_admin_screen.png",
        "caption": (
            "운영자는 관제탑의 보상 마켓 탭에서 발급 상태, 비즈머니/준비금, 최근 발급 내역, "
            "수동 재확인 업무를 관리합니다. 사용자 교환 화면과 운영 화면을 분리해 안정적으로 운영합니다."
        ),
    },
]

OVERVIEW_ROWS = [
    ("서비스명", "해빛스쿨(Habit School)"),
    ("운영사", "(주)공감케어"),
    ("서비스 성격", "건강 습관 형성 보상 앱"),
    ("기프티콘 교환 자산", "앱 내 포인트"),
    ("현재 운영 정책", "포인트 차감 후 Giftishow API 연동 쿠폰 발급"),
    ("쿠폰 수령 방식", "앱 내 쿠폰 보관함 우선 노출"),
    ("관리 방식", "관제탑에서 발급 현황 및 수동 재확인 운영"),
    ("초기 대표 상품", "메가MGC커피 (ICE)아메리카노 모바일쿠폰"),
]

FLOW_BULLETS = [
    "회원이 건강 챌린지와 앱 활동을 통해 포인트를 적립합니다.",
    "회원이 해빛 마켓에서 원하는 보상 상품을 선택합니다.",
    "앱 서버가 포인트 차감 후 Giftishow API로 쿠폰 발급을 요청합니다.",
    "발급된 쿠폰은 앱 내 쿠폰 보관함에서 바코드 이미지 또는 PIN으로 확인합니다.",
    "운영자는 관제탑에서 발급 상태, 예산, 수동 재확인 건을 관리합니다.",
]

LIVE_NOTES = [
    "포인트 기반 보상몰 운영 구조이며, 현재 기프티콘 교환은 HBT가 아닌 앱 내 포인트로 처리합니다.",
    "쿠폰은 기본적으로 앱 내 보관함에서 확인하도록 설계했고, 운영 예외 시에만 별도 재전송 정책을 검토합니다.",
    "상용 초기 운영 상품은 소액 보상형 상품부터 시작하며, 메가MGC커피 상품을 1차 대상으로 검토하고 있습니다.",
    "사용자 연락처는 공급사 발급 조건과 고객 지원을 위한 최소 범위에서만 사용합니다.",
]


def add_paragraph(document, text, *, size=11, bold=False, space_after=8, alignment=None):
    paragraph = document.add_paragraph()
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph


def add_bullets(document, items, *, font_size=10.5):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.size = Pt(font_size)
        paragraph.paragraph_format.space_after = Pt(3)


def add_heading(document, text, *, size=14):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    paragraph.paragraph_format.space_after = Pt(8)
    return paragraph


def build_markdown_source():
    lines = [
        "# 해빛스쿨 Giftishow 상용 Key 승인 제출 문서",
        "",
        "## 서비스 개요",
        "",
    ]
    for label, value in OVERVIEW_ROWS:
        lines.append(f"- {label}: {value}")
    lines.extend(
        [
            "",
            "## 서비스 흐름",
            "",
        ]
    )
    for item in FLOW_BULLETS:
        lines.append(f"- {item}")
    lines.extend(
        [
            "",
            "## 운영 메모",
            "",
        ]
    )
    for note in LIVE_NOTES:
        lines.append(f"- {note}")
    SOURCE_NOTE.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    build_markdown_source()

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    add_paragraph(
        document,
        "해빛스쿨 Giftishow 상용 Key 승인 제출 문서",
        size=20,
        bold=True,
        space_after=2,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        document,
        "실서비스 적용 화면 및 서비스 소개",
        size=11,
        space_after=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    add_paragraph(
        document,
        "해빛스쿨은 건강 습관 형성 과정을 기록하고 보상으로 연결하는 모바일 서비스입니다. "
        "회원은 앱 내 활동을 통해 포인트를 적립하고, 해빛 마켓에서 Giftishow 기반 모바일 쿠폰으로 교환할 수 있습니다. "
        "발급된 쿠폰은 앱 내 쿠폰 보관함에서 직접 확인하도록 구현했습니다.",
        size=11,
        space_after=8,
    )
    add_paragraph(
        document,
        "아래 자료는 현재 서비스에 적용된 사용자 자산 탭, 쿠폰 보관함, 관리자 관제탑의 화면 구성을 정리한 제출용 문서입니다.",
        size=11,
        space_after=10,
    )

    add_heading(document, "1. 서비스 개요")
    for label, value in OVERVIEW_ROWS:
        paragraph = document.add_paragraph()
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.size = Pt(10.5)
        value_run = paragraph.add_run(value)
        value_run.font.size = Pt(10.5)
        paragraph.paragraph_format.space_after = Pt(4)

    add_heading(document, "2. 서비스 흐름")
    add_bullets(document, FLOW_BULLETS)

    add_heading(document, "3. 운영 메모")
    add_bullets(document, LIVE_NOTES)

    add_paragraph(
        document,
        "참고: 기프티콘 교환은 현재 앱 내 포인트 기준으로 운영하며, HBT는 별도 자산 시스템으로 유지하고 있습니다.",
        size=10.5,
        bold=True,
        space_after=12,
    )

    for index, screen in enumerate(SCREENS):
        if index == 0:
            document.add_page_break()
        else:
            document.add_section(WD_SECTION.NEW_PAGE)

        add_heading(document, screen["heading"], size=15)
        image_path = screen["path"]
        if not image_path.exists():
            raise FileNotFoundError(f"Missing screenshot: {image_path}")

        document.add_picture(str(image_path), width=Inches(6.45))
        picture_paragraph = document.paragraphs[-1]
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.paragraph_format.space_after = Pt(10)

        add_paragraph(document, screen["caption"], size=10.5, space_after=10)

    add_paragraph(
        document,
        "해빛스쿨 | Habit School Reward Market | 2026-04-24",
        size=9,
        space_after=0,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    document.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
